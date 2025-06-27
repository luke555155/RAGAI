from uuid import uuid4
from typing import List
from time import time, sleep
import io
import re
import json
import logging
from datetime import datetime

from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse
from fastapi.exceptions import RequestValidationError
from pydantic import BaseModel
from langchain.text_splitter import RecursiveCharacterTextSplitter
import PyPDF2
import docx
import requests
from dotenv import load_dotenv
import os

LOG_FORMAT = "%(asctime)s [%(name)s] %(levelname)s %(message)s"
os.makedirs("logs", exist_ok=True)
logging.basicConfig(
    level=logging.INFO,
    format=LOG_FORMAT,
    handlers=[
        logging.FileHandler("logs/app.log"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger(__name__)
load_dotenv()

from .services.qdrant_client import QdrantClient
from .upload import (
    upload_document,
    load_doc_index,
    save_doc_index,
    remove_doc_entry,
)
from .reranker import rerank_passages, cosine_similarity

app = FastAPI(docs_url="/docs", openapi_url="/api/openapi.json")

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434").rstrip("/")
OLLAMA_EMBEDDING_URL = f"{OLLAMA_BASE_URL}/api/embeddings"
OLLAMA_GENERATE_URL = f"{OLLAMA_BASE_URL}/api/generate"
QDRANT_HOST = os.getenv("QDRANT_HOST", "localhost")
QDRANT_PORT = os.getenv("QDRANT_PORT", "6333")
QDRANT_URL = f"http://{QDRANT_HOST}:{QDRANT_PORT}"
COLLECTION_NAME = "documents"
OLLAMA_MODEL = os.getenv("EMBEDDING_MODEL", "nomic-embed-text")
OLLAMA_LLM_MODEL = os.getenv("LLM_MODEL", "llama2")
RERANK_MODE = os.getenv("RERANK_MODE", "llm")
SEARCH_LIMIT = int(os.getenv("SEARCH_LIMIT", "8"))

MAX_CONTEXT_TOKENS = int(os.getenv("MAX_CONTEXT_TOKENS", "1500"))
CACHE_TTL = int(os.getenv("ASK_CACHE_TTL", "300"))
ASK_CACHE: dict = {}

qdrant_client = QdrantClient(QDRANT_URL, COLLECTION_NAME)


def ensure_paths() -> None:
    os.makedirs("logs", exist_ok=True)
    os.makedirs("ollama", exist_ok=True)
    if not os.path.exists("doc_index.json"):
        with open("doc_index.json", "w", encoding="utf-8") as f:
            json.dump({}, f)
    # create example document
    example_path = "example.docx"
    if not os.path.exists(example_path):
        try:
            document = docx.Document()
            document.add_paragraph("這是一個範例文件，供測試上傳功能。")
            document.save(example_path)
        except Exception as exc:
            logger.warning("Failed to create example document: %s", exc)


def check_ollama_ready() -> None:
    payload = {"name": OLLAMA_LLM_MODEL}
    while True:
        try:
            resp = requests.post(f"{OLLAMA_BASE_URL}/api/show", json=payload)
            if resp.status_code == 200:
                break
            logger.warning("Ollama model not ready, retrying...")
        except Exception as exc:
            logger.warning("Ollama connection failed: %s", exc)
        sleep(3)


@app.on_event("startup")
async def startup_event() -> None:
    ensure_paths()
    check_ollama_ready()


@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc: HTTPException):
    logger.error("HTTP error on %s: %s", request.url.path, exc.detail)
    return JSONResponse(status_code=exc.status_code, content={"error": exc.detail})


@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request, exc: RequestValidationError):
    logger.error("Validation error on %s: %s", request.url.path, exc)
    return JSONResponse(status_code=422, content={"error": str(exc)})


@app.exception_handler(Exception)
async def unhandled_exception_handler(request, exc: Exception):
    logger.exception("Unhandled error on %s", request.url.path)
    return JSONResponse(status_code=500, content={"error": "Internal server error"})




def read_pdf(file_bytes: bytes) -> str:
    """讀取 PDF 檔案並回傳文字內容"""
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text


def read_docx(file_bytes: bytes) -> str:
    """讀取 Word 檔案並回傳文字內容"""
    document = docx.Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in document.paragraphs)


def split_text(text: str) -> List[str]:
    """將文本依長度切分為數個區塊，中文以句號等標點為界"""
    if re.search(r"[\u4e00-\u9fff]", text):
        sentences = re.split(r"(?<=[。！？])", text)
        sentences = [s.strip() for s in sentences if s.strip()]
        chunks = []
        current = ""
        for sent in sentences:
            if len(current) + len(sent) > 500:
                chunks.append(current)
                current = sent
            else:
                current += sent
        if current:
            chunks.append(current)
        return chunks
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    return splitter.split_text(text)


def get_embedding(text: str) -> List[float]:
    """向 Ollama 取得文字向量，失敗時回傳 500"""
    try:
        resp = requests.post(
            OLLAMA_EMBEDDING_URL,
            json={"model": OLLAMA_MODEL, "prompt": text},
        )
        resp.raise_for_status()
        embedding = resp.json().get("embedding")
        if embedding is None:
            raise ValueError("No embedding returned")
        return embedding
    except Exception as exc:
        logger.exception("Embedding service error: %s", exc)
        raise HTTPException(status_code=500, detail=f"Embedding service error: {exc}") from exc


def count_tokens(text: str) -> int:
    """Rudimentary token count based on whitespace splitting"""
    return len(text.split())


class StatusResponse(BaseModel):
    ollama: bool
    qdrant: bool
    embedding_model: bool


@app.get("/api/status", tags=["system"], response_model=StatusResponse, description="Return health status of backend services")
async def get_status() -> StatusResponse:
    ollama_ok = False
    qdrant_ok = False
    embed_ok = False
    try:
        resp = requests.post(f"{OLLAMA_BASE_URL}/api/show", json={"name": OLLAMA_LLM_MODEL}, timeout=5)
        ollama_ok = resp.status_code == 200
        if ollama_ok:
            test_resp = requests.post(OLLAMA_EMBEDDING_URL, json={"model": OLLAMA_MODEL, "prompt": "hello"}, timeout=5)
            embed_ok = test_resp.status_code == 200
    except Exception as exc:
        logger.warning("Ollama check failed: %s", exc)
    try:
        qdrant_client.ensure_collection()
        qdrant_ok = True
    except Exception as exc:
        logger.warning("Qdrant check failed: %s", exc)
    return StatusResponse(ollama=ollama_ok, qdrant=qdrant_ok, embedding_model=embed_ok)






class UploadResponse(BaseModel):
    document_id: str
    segments_uploaded: int
    summary: str


@app.post(
    "/api/upload",
    tags=["documents"],
    response_model=UploadResponse,
    description="Upload a document and store its embeddings",
)
async def upload(file: UploadFile = File(...), tags: str = Form("")) -> UploadResponse:
    """上傳檔案並分割後寫入 Qdrant"""
    content = await file.read()
    if file.content_type == "application/pdf":
        text = read_pdf(content)
    elif file.content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
        text = read_docx(content)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    chunks = split_text(text)
    document_id = str(uuid4())
    tag_list = [t.strip() for t in tags.split(',') if t.strip()]
    summary = upload_document(document_id, chunks, file.filename, tag_list)
    return UploadResponse(document_id=document_id, segments_uploaded=len(chunks), summary=summary)


class AskRequest(BaseModel):
    question: str
    document_id: str | None = None
    document_ids: List[str] | None = None
    style: str | None = None
    prompt_template: str | None = None
    rerank_mode: str | None = None
    include_filtered: bool | None = True


class DeleteDocRequest(BaseModel):
    document_id: str


class RankTestRequest(BaseModel):
    question: str
    passages: List[str]
    mode: str | None = None


class ManualAskRequest(BaseModel):
    question: str
    segments: List[str]
    style: str | None = None
    prompt_template: str | None = None


class DocumentInfo(BaseModel):
    document_id: str
    file_name: str | None = None
    upload_time: str | None = None
    summary: str | None = None
    tags: List[str] | None = None


class ListDocsResponse(BaseModel):
    documents: List[DocumentInfo]


class DocumentSegmentsResponse(BaseModel):
    document_id: str
    segments: List[dict]


class ResummarizeResponse(BaseModel):
    document_id: str
    summary: str


class RankTestResponse(BaseModel):
    answer: str
    references: List[dict]


class AskResponse(BaseModel):
    answer: str
    references: List[dict]
    elapsed: float


class AskLogResponse(BaseModel):
    logs: List[dict]


class ManualAskResponse(BaseModel):
    answer: str
    references: List[dict]
    elapsed: float


@app.get(
    "/api/docs",
    tags=["documents"],
    response_model=ListDocsResponse,
    description="List all uploaded documents",
)
async def list_docs() -> ListDocsResponse:
    """回傳所有文件資訊"""
    try:
        docs = qdrant_client.list_documents()
        summaries = load_doc_index()
        for doc in docs:
            doc_id = doc.get("document_id")
            if doc_id in summaries:
                entry = summaries[doc_id]
                if isinstance(entry, dict):
                    doc["summary"] = entry.get("summary")
                    doc["tags"] = entry.get("tags", [])
                else:
                    # backward compatibility if index contains only summary
                    doc["summary"] = entry
                    doc["tags"] = []
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("List docs error: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    return ListDocsResponse(documents=[DocumentInfo(**d) for d in docs])


@app.get(
    "/api/docs/{document_id}",
    tags=["documents"],
    response_model=DocumentSegmentsResponse,
    description="Get all segments of a document",
)
async def get_document(document_id: str) -> DocumentSegmentsResponse:
    """取得指定文件的段落內容"""
    try:
        segments = qdrant_client.get_segments_by_doc(document_id)
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Get document error: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    return DocumentSegmentsResponse(document_id=document_id, segments=segments)


class DeleteResponse(BaseModel):
    status: str


@app.delete(
    "/api/docs",
    tags=["documents"],
    response_model=DeleteResponse,
    description="Delete a document and its vectors",
)
async def delete_document(req: DeleteDocRequest) -> DeleteResponse:
    """刪除指定文件的所有資料"""
    try:
        qdrant_client.delete_by_document(req.document_id)
        remove_doc_entry(req.document_id)
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Delete document error: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    return DeleteResponse(status="deleted")


@app.post(
    "/api/docs/{document_id}/resummarize",
    tags=["documents"],
    response_model=ResummarizeResponse,
    description="Regenerate summary for a document",
)
async def resummarize_document(document_id: str) -> ResummarizeResponse:
    """重新取得段落並產生摘要"""
    try:
        segments = qdrant_client.get_segments_by_doc(document_id)
        if not segments:
            raise HTTPException(status_code=404, detail="Document not found")
        context = "\n".join(s.get("text", "") for s in segments)
        resp = requests.post(
            OLLAMA_GENERATE_URL,
            json={"model": OLLAMA_LLM_MODEL, "prompt": f"請根據以下內容提供一段摘要：\n{context}"},
        )
        resp.raise_for_status()
        summary = resp.json().get("response", "")
        index = load_doc_index()
        entry = index.get(document_id, {})
        entry["summary"] = summary
        if "tags" not in entry:
            entry["tags"] = []
        index[document_id] = entry
        save_doc_index(index)
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Resummarize error: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    return ResummarizeResponse(document_id=document_id, summary=summary)


@app.post(
    "/api/rank_test",
    tags=["debug"],
    response_model=RankTestResponse,
    description="Return ranking results for provided passages",
)
async def rank_test(req: RankTestRequest) -> RankTestResponse:
    """Return ranking results for provided passages"""
    try:
        q_emb = get_embedding(req.question)
        passages = []
        for idx, text in enumerate(req.passages):
            emb = get_embedding(text)
            score = cosine_similarity(q_emb, emb)
            passages.append({"id": str(idx), "payload": {"text": text}, "vector": emb, "score": score})
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Rank test error: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    ranked = rerank_passages(req.question, passages, mode=req.mode or RERANK_MODE)
    refs = [
        {
            "text": p["payload"].get("text", ""),
            "rank": p.get("rank"),
            "score": p.get("score"),
            "filtered": p.get("filtered"),
        }
        for p in ranked
    ]
    return RankTestResponse(answer="", references=refs)




@app.post(
    "/api/ask",
    tags=["ask"],
    response_model=AskResponse,
    description="Ask question based on uploaded documents",
)
async def ask(req: AskRequest) -> AskResponse:
    """根據問題向 Qdrant 取得相關段落並呼叫 LLM 回答"""
    cache_key = (
        req.question,
        tuple(req.document_ids or ([req.document_id] if req.document_id else [])),
        req.style,
        req.prompt_template,
        req.rerank_mode,
        req.include_filtered,
    )
    cached = ASK_CACHE.get(cache_key)
    if cached and time() - cached["ts"] < CACHE_TTL:
        return cached["data"]

    start_time = time()
    q_embedding = get_embedding(req.question)
    try:
        results = qdrant_client.search(
            q_embedding,
            limit=SEARCH_LIMIT,
            document_id=req.document_id,
            document_ids=req.document_ids,
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Search Qdrant error: %s", exc)
        raise HTTPException(status_code=500, detail=f"Qdrant search error: {exc}") from exc

    ranked_all = rerank_passages(req.question, results, mode=req.rerank_mode or RERANK_MODE)
    for idx, r in enumerate(ranked_all, 1):
        if (req.rerank_mode or RERANK_MODE) == "llm" and idx > 5:
            r["filtered"] = True
    context_parts = []
    used_results = []
    tokens = 0
    for r in ranked_all:
        if r.get("filtered") and not req.include_filtered:
            continue
        text = r["payload"].get("text", "")
        t = count_tokens(text)
        if tokens + t > MAX_CONTEXT_TOKENS:
            break
        tokens += t
        context_parts.append(text)
        used_results.append(r)
    context = "\n".join(context_parts)

    template = (
        req.prompt_template
        or "Answer the question based on the context below.\n{style}\n\nContext:\n{context}\n\nQuestion: {question}\nAnswer:"
    )
    prompt = template.format(
        context=context,
        question=req.question,
        style=req.style or "",
    )
    try:
        resp = requests.post(
            OLLAMA_GENERATE_URL,
            json={"model": OLLAMA_LLM_MODEL, "prompt": prompt},
        )
        resp.raise_for_status()
        answer = resp.json().get("response", "")
    except Exception as exc:
        logger.exception("Ollama service error: %s", exc)
        raise HTTPException(status_code=500, detail=f"Ollama service error: {exc}") from exc
    references = [
        {
            "text": r["payload"].get("text", ""),
            "chunk_index": r["payload"].get("chunk_index"),
            "score": r.get("score"),
            "rank": r.get("rank"),
            "filtered": r.get("filtered"),
        }
        for r in used_results
    ]
    elapsed = time() - start_time
    data = AskResponse(answer=answer, references=references, elapsed=elapsed)
    try:
        os.makedirs("logs", exist_ok=True)
        with open("logs/ask_log.jsonl", "a", encoding="utf-8") as f:
            json.dump(
                {
                    "timestamp": datetime.utcnow().isoformat(),
                    "question": req.question,
                    "document_ids": req.document_ids
                    or ([req.document_id] if req.document_id else []),
                    "style": req.style,
                    "rerank_mode": req.rerank_mode or RERANK_MODE,
                    "answer": answer,
                    "references": references,
                    "elapsed": elapsed,
                },
                f,
                ensure_ascii=False,
            )
            f.write("\n")
    except Exception as exc:
        logger.warning("Failed to log ask: %s", exc)
    ASK_CACHE[cache_key] = {"ts": time(), "data": data}
    return data


@app.get(
    "/api/ask_log",
    tags=["ask"],
    response_model=AskLogResponse,
    description="Retrieve ask history logs",
)
async def get_ask_log() -> AskLogResponse:
    """Return logged ask records"""
    try:
        entries = []
        if os.path.exists("logs/ask_log.jsonl"):
            with open("logs/ask_log.jsonl", "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        entries.append(json.loads(line))
                    except json.JSONDecodeError:
                        continue
        return AskLogResponse(logs=entries[::-1])
    except Exception as exc:
        logger.exception("Read ask log error: %s", exc)
        raise HTTPException(status_code=500, detail="Failed to read log") from exc


@app.post(
    "/api/manual-ask",
    tags=["ask"],
    response_model=ManualAskResponse,
    description="Ask a question using manually provided context",
)
async def manual_ask(req: ManualAskRequest) -> ManualAskResponse:
    """Answer a question using manually provided context"""
    start_time = time()
    context = "\n".join(req.segments)
    template = (
        req.prompt_template
        or "Answer the question based on the context below.\n{style}\n\nContext:\n{context}\n\nQuestion: {question}\nAnswer:"
    )
    prompt = template.format(context=context, question=req.question, style=req.style or "")
    try:
        resp = requests.post(
            OLLAMA_GENERATE_URL,
            json={"model": OLLAMA_LLM_MODEL, "prompt": prompt},
        )
        resp.raise_for_status()
        answer = resp.json().get("response", "")
    except Exception as exc:
        logger.exception("Ollama service error: %s", exc)
        raise HTTPException(status_code=500, detail=f"Ollama service error: {exc}") from exc
    elapsed = time() - start_time
    refs = [{"text": s} for s in req.segments]
    return ManualAskResponse(answer=answer, references=refs, elapsed=elapsed)
